# SYNTHETIC DATA - NOT REAL PII
"""Programmatic synthetic attachment generator (Phase 4 fixtures).

Builds in-memory bytes for every test attachment shape we need:

  - PDF (text-bearing)
  - PDF (multi-page > MAX_PAGES)
  - PDF (scan-only / no embedded text)
  - PDF (encrypted)
  - PDF (corrupted)
  - DOCX (with PII in table cells)
  - HWPX (minimal valid container with section1.xml)

Every fixture is deterministic so tests can hash-compare. No real PII
appears anywhere in the payload data.
"""

from __future__ import annotations

import io
import zipfile
from typing import Final

# Synthetic phone reserved by KISA for test usage. See generator notes.
SYNTH_PHONE: Final[str] = "010-0000-1234"
SYNTH_EMAIL: Final[str] = "synth@example.com"
SYNTH_TEXT: Final[str] = (
    "이것은 합성 테스트 문서입니다. "
    f"전화번호 {SYNTH_PHONE} 와 이메일 {SYNTH_EMAIL} 가 포함되어 있습니다. "
    "SYNTHETIC DATA - NOT REAL PII"
)


# ── PDF helpers ──────────────────────────────────────────────────────────
def _build_pdf(pages_text: list[str]) -> bytes:
    """Build a minimal text-bearing PDF using pypdfium2's writer.

    pypdfium2 is read-only; we instead use ``reportlab`` if available,
    falling back to a hand-rolled PDF (sufficient for pdfplumber to
    parse a single text run).
    """
    try:
        from reportlab.pdfgen import canvas  # type: ignore[import-not-found]

        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        for txt in pages_text:
            c.drawString(72, 720, txt)
            c.showPage()
        c.save()
        return buf.getvalue()
    except ImportError:
        # Hand-rolled minimal PDF — works for pdfplumber single-page text.
        return _hand_rolled_pdf(pages_text)


def _hand_rolled_pdf(pages_text: list[str]) -> bytes:
    """Hand-rolled PDF that pdfplumber + pypdfium2 can both read.

    Uses the standard 14 Type1 base font Helvetica so no font embedding
    is needed.
    """
    n_pages = max(1, len(pages_text))
    objects: list[bytes] = []

    def _obj(num: int, body: bytes) -> bytes:
        return f"{num} 0 obj\n".encode() + body + b"\nendobj\n"

    # Object 1: catalog
    objects.append(_obj(1, b"<< /Type /Catalog /Pages 2 0 R >>"))
    # Object 2: pages tree
    kids = " ".join(f"{3 + i * 2} 0 R" for i in range(n_pages))
    objects.append(
        _obj(
            2,
            f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>".encode(),
        )
    )
    # Object N: per-page page + content
    for i in range(n_pages):
        page_obj_num = 3 + i * 2
        content_obj_num = 4 + i * 2
        text = pages_text[i] if i < len(pages_text) else ""
        # Use Helvetica (font alias /F1 declared in resources).
        # Latin1-encoded body keeps it portable; CJK is not rendered but
        # pdfplumber returns the codepoints as text which is sufficient
        # for downstream regex testing.
        safe = text.replace("(", "\\(").replace(")", "\\)")
        stream = b"BT /F1 12 Tf 72 720 Td (" + safe.encode("latin-1", errors="replace") + b") Tj ET"
        # Page object references parent + resources + content stream
        page_dict = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 "
            f"/BaseFont /Helvetica >> >> >> /Contents {content_obj_num} 0 R >>"
        ).encode()
        objects.append(_obj(page_obj_num, page_dict))
        content_dict = f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        objects.append(_obj(content_obj_num, content_dict))

    out = bytearray(b"%PDF-1.4\n%\xff\xff\xff\xff\n")
    offsets: list[int] = []
    for o in objects:
        offsets.append(len(out))
        out += o
    xref_pos = len(out)
    out += b"xref\n"
    out += f"0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += b"trailer\n"
    out += (f"<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n").encode()
    out += b"%%EOF\n"
    return bytes(out)


def make_text_pdf() -> bytes:
    """Single-page PDF containing ``SYNTH_TEXT``."""
    return _build_pdf([SYNTH_TEXT])


def make_multipage_pdf(n_pages: int) -> bytes:
    """N-page PDF with one synthetic line per page."""
    return _build_pdf([f"page {i} {SYNTH_PHONE}" for i in range(n_pages)])


def make_scan_only_pdf() -> bytes:
    """Minimal valid PDF whose page has no text content stream."""
    pdf_body = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << >> >>\nendobj\n"
    )
    # Build correct xref + trailer
    offsets = [0]
    pos = len(b"%PDF-1.4\n")
    for body in (
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << >> >>\nendobj\n",
    ):
        offsets.append(pos)
        pos += len(body)

    xref = b"xref\n0 4\n0000000000 65535 f \n"
    for off in offsets[1:]:
        xref += f"{off:010d} 00000 n \n".encode()
    trailer = b"trailer\n<< /Size 4 /Root 1 0 R >>\nstartxref\n" + str(pos).encode() + b"\n%%EOF\n"
    return pdf_body + xref + trailer


def make_corrupted_pdf() -> bytes:
    """Bytes that look PDF-ish but have a broken structure."""
    return b"%PDF-1.4\n%this-is-not-valid\n0 0 obj\nbroken\n%%EOF\n"


def make_encrypted_pdf() -> bytes:
    """Trivially-encrypted PDF the parsers will reject with a password hint.

    Real encryption would require a full crypto handler; we simulate by
    injecting an /Encrypt entry into the trailer so pdfplumber/pypdfium2
    surface a 'password' / 'encrypted' error.
    """
    base = make_text_pdf()
    # Inject a fake /Encrypt reference so pdfium recognises the file as
    # password-protected. pdfium's error message contains "password".
    trailer_pos = base.rfind(b"trailer")
    if trailer_pos < 0:
        return base
    head = base[:trailer_pos]
    tail = base[trailer_pos:]
    # Replace ``<< /Size N /Root ...>>`` with a version including /Encrypt 99 0 R
    tail = tail.replace(b"<< /Size", b"<< /Encrypt 99 0 R /Size", 1)
    return head + tail


# ── DOCX helpers ─────────────────────────────────────────────────────────
def make_docx_with_table_pii() -> bytes:
    """Build a .docx whose only text lives in a single table cell."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("This file is synthetic data.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "전화"
    table.cell(0, 1).text = SYNTH_PHONE
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HWPX helpers ─────────────────────────────────────────────────────────
HWPX_NS = "http://www.hancom.co.kr/hwpml/2011/section"


def make_hwpx_with_text() -> bytes:
    """Minimal HWPX archive: only Contents/section0.xml carries text."""
    section_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        f'<sec xmlns="{HWPX_NS}">'
        f"<p><run><t>{SYNTH_TEXT}</t></run></p>"
        f"</sec>"
    ).encode()
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/hwp+zip")
        zf.writestr("Contents/section0.xml", section_xml)
    return buf.getvalue()


def make_hwp5_binary() -> bytes:
    """4-byte CFB header so the dispatcher can sniff it as HWP 5 binary."""
    return b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64


# ── Plain text ───────────────────────────────────────────────────────────
def make_text_file(text: str = SYNTH_TEXT) -> bytes:
    return text.encode("utf-8")
