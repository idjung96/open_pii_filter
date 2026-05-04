"""DOCX text extraction (Phase 4, T4.3).

Uses ``python-docx`` to extract paragraph text and table cell text.
Password-protected (encrypted) DOCX files have an OOXML container that
``python-docx`` can't open; we map that to REQ-4051 so the caller sees
the right error code.
"""

from __future__ import annotations

import asyncio
import io
import zipfile
from typing import TYPE_CHECKING

from app.extractors.fetcher import ExtractionError

if TYPE_CHECKING:
    pass


def _is_encrypted_ooxml(data: bytes) -> bool:
    """Detect MS Office DRM/password protection (CDF compound file format).

    Encrypted .docx files are wrapped in an MS-CFB container (D0 CF 11 E0)
    instead of the normal ZIP container.
    """
    return data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _extract_sync(data: bytes, filename: str) -> str:
    """Synchronous core run on a worker thread."""
    import docx

    if _is_encrypted_ooxml(data):
        raise ExtractionError("REQ-4051", filename=filename, detail="password-protected")

    try:
        document = docx.Document(io.BytesIO(data))
    except zipfile.BadZipFile as e:
        raise ExtractionError("REQ-4042", filename=filename, detail=f"bad zip: {e}") from e
    except Exception as e:
        raise ExtractionError("REQ-4042", filename=filename, detail=str(e)) from e

    parts: list[str] = []
    try:
        for para in document.paragraphs:
            if para.text:
                parts.append(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
    except Exception as e:
        raise ExtractionError("REQ-4042", filename=filename, detail=f"walk failed: {e}") from e

    return "\n".join(parts)


async def extract_docx(data: bytes, filename: str) -> str:
    """Extract concatenated text from a .docx file."""
    return await asyncio.to_thread(_extract_sync, data, filename)
